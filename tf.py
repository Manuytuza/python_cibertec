"""
import random
import matplotlib.pyplot as plt
import pandas as pd

#lista de los 50 valores de moneda
moneda_valor = []
fecha_moneda =[]
#50 nombres de moneda
mo_name = ["sol", "dolar", "euro", "libra", "yen", "franco", "peso", "real", "lira", "rupia", "won", "dinar", "dirham", "shekel", "krone", "koruna", "forint", "leu", "baht", "dong", "rand", "peso argentino", "peso chileno", "peso colombiano", "peso mexicano", "peso peruano", "peso uruguayo", "peso venezolano", "peso boliviano", "peso paraguayo", "peso guarani", "peso hondureño", "peso salvadoreño", "peso nicaragüense", "peso costarricense", "peso panameño", "peso dominicano", "peso cubano", "peso filipino", "peso guatemalteco", "peso beliceño", "peso jamaicano", "peso trinitense", "peso bahameño", "peso barbadeño", "peso grenadino", "peso sanvicentino", "peso santa luciano", "peso dominicano", "peso antillano"] 

#bucle for generador de monedas con random
for _ in range(50):
    #genera valores de moneda
    m_v = round(random.uniform(1,5),2) 
    moneda_valor.append(m_v)
    #genera fecha de lanzamiento de moneda
    fecha_dia = random.randint(1,30)
    fecha_moneda.append(fecha_dia)

df= pd.DataFrame({
     "Name": mo_name,
     "Date" : fecha_moneda,
     "Value" : moneda_valor
 })

df = df.sort_values(by="Value")
print(df)


plt.plot(df["Name"], df["Value"], label = "Valor por Moneda")
plt.plot(df["Name"], df["Date"], label = "Fecha de Lanzamiento")

plt.xlabel("Nombre")
plt.ylabel("Valor")

plt.xticks(rotation=90)
plt.legend()
plt.show()

#---------------------------------------
import random
import pandas as pd
import matplotlib.pyplot as plt

moneda = ["Cara", "Sello"]

# 50 lanzamientos
lanzamientos = [random.choice(moneda) for _ in range(50)]

# crear DataFrame
df = pd.DataFrame({
    "Resultado": lanzamientos
})

print(df)

#contar resultados
conteo = df["Resultado"].value_counts()

print("\nConteo:")
print(conteo)

#crear grafico 
conteo.plot(kind="bar", title="Resultados de 50 lanzamientos de moneda")

plt.xlabel("Resultado")
plt.ylabel("Cantidad")
plt.xticks(rotation=0)

plt.tight_layout()
plt.show()

#ejercicoo moneda trampoza ------------------------------------

lanzamientos = random.choices(
    ["Cara", "Sello"],
    weights=[0.7, 0.3],
    k=50
)
print(lanzamientos)

df1 = pd.DataFrame({
    "Resultado2": lanzamientos}
    )
print(df1)

conteo2 = df1["Resultado2"].value_counts()

plt.pie(
    conteo2.values,
    labels=conteo.index,
    autopct="%1.1f%%",
    startangle=90,
    colors=["#66b3ff", "#ffcf99"],
    wedgeprops={"edgecolor": "black"}
)

plt.title("Resultados de la moneda tramposa")

plt.tight_layout()
plt.show()

"""
import pandas as pd

# =========================
# crear calendario 2020
# =========================

# generar fechas
fechas = pd.date_range(start="2020-01-01", end="2020-12-31")

# crear dataframe
df_calendario = pd.DataFrame({
    "Date": fechas,
    "Day_of_week": fechas.day_name(),
    "Week_number": fechas.isocalendar().week,
    "Month": fechas.month_name(),
    "Quarter": fechas.quarter
})

# =========================
# exportar a excel
# =========================

df_calendario.to_excel("Calendario_2020.xlsx", index=False)

print("excel generado correctamente")

# =========================
# exportar a word
# =========================

from docx import Document

# crear documento
doc = Document()
doc.add_heading("Calendario 2020", level=1)

# crear tabla
tabla = doc.add_table(rows=1, cols=len(df_calendario.columns))
tabla.style = "Table Grid"

# encabezados
hdr_cells = tabla.rows[0].cells
for i, col in enumerate(df_calendario.columns):
    hdr_cells[i].text = col

# datos
for _, row in df_calendario.iterrows():
    fila = tabla.add_row().cells
    for i, col in enumerate(df_calendario.columns):
        fila[i].text = str(row[col])

# guardar word
doc.save("Calendario_2020.docx")

print("word generado correctamente")